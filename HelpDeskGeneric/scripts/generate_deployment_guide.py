"""
汎用ヘルプデスク 導入手順書 生成スクリプト

solution-docs スキルのテンプレートに基づいて Word docx を生成する。
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# ============================================================
# 設定値
# ============================================================
APP_NAME = "汎用ヘルプデスク"
SOLUTION_VERSION = "1.1.0.0"
DOC_DATE = "2026年4月"
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, f"{APP_NAME}_ソリューション導入手順書.docx")

# ソリューション構成テーブル
COMPONENTS = [
    ("テーブル", "5", "チケット、カテゴリ、ナレッジ記事、BPF×2"),
    ("モデル駆動型アプリ", "1", "汎用ヘルプデスク"),
    ("Power Automate フロー", "3", "チケット起票通知、クローズ時ナレッジ化、エージェント連携"),
    ("BPF", "2", "チケット管理プロセス、ナレッジ記事管理プロセス"),
    ("Copilot Studio Bot", "2", "Ask Copilot（担当者用）、Ask Specialist（ユーザー用）"),
    ("システムチャート", "9", "ステータス別、優先度別、カテゴリ別、起票推移等"),
    ("接続参照", "3", "Dataverse、Copilot Studio、Office 365 Outlook"),
    ("Web Resource", "3", "SVG アイコン（チケット、ナレッジ、カテゴリ）"),
]

# 必要な条件テーブル
PREREQUISITES = [
    ("Power Platform 環境", "Dataverse が有効であること"),
    ("ライセンス", "Power Apps Premium ライセンス（開発者・利用者）"),
    ("Copilot Studio", "Microsoft Copilot Studio ライセンス（エージェント機能使用時）"),
    ("セキュリティロール", "System Administrator（インストール実施者）"),
    ("ブラウザ", "Microsoft Edge / Google Chrome（最新版）"),
]

# 接続テーブル
CONNECTIONS = [
    ("Microsoft Dataverse", "テーブルの読み書き"),
    ("Microsoft Copilot Studio", "エージェント連携"),
    ("Office 365 Outlook", "通知メール送信"),
]

# フローテーブル
FLOWS = [
    ("チケットが起票された場合に呼び出しされる", "チケット作成時に通知", "有効化 + 通知先メールアドレスを確認"),
    ("チケットがクローズした場合にナレッジ化", "クローズ時にナレッジ記事を自動作成", "有効化"),
    ("エージェントフロー_チケットを起票", "エージェントからチケットを自動起票", "必要に応じて有効化"),
]

# カテゴリ推奨構成
CATEGORIES = [
    ("IT インフラ", "ネットワーク接続、セキュリティ"),
    ("アプリケーション", "Microsoft 365、業務システム"),
    ("アカウント管理", "パスワードリセット"),
    ("ハードウェア", "PC 故障・交換"),
    ("その他", "—"),
]

# ナビゲーション
NAVIGATION = [
    ("チケット管理", "チケット", "問い合わせチケットの一覧・管理"),
    ("ナレッジ", "ナレッジ記事", "FAQ やノウハウ記事の管理"),
    ("マスタデータ", "カテゴリ", "チケット・記事の分類カテゴリ"),
]

# ステータス遷移
STATUSES = [
    ("新規", "起票直後", "対応中 / 保留"),
    ("対応中", "担当者が対応中", "解決済 / 保留 / エスカレーション"),
    ("保留", "外部回答待ち等", "対応中"),
    ("エスカレーション", "上位者にエスカレーション", "対応中 / 解決済"),
    ("解決済", "回答完了", "クローズ / 対応中（再オープン）"),
    ("クローズ", "完了", "—"),
]

# チケットテーブル列定義
TICKET_COLUMNS = [
    ("タイトル", "new_title", "テキスト", "問い合わせの件名（必須）"),
    ("チケット番号", "new_ticketnumber", "テキスト", "自動採番（読取専用）"),
    ("ステータス", "new_ticketstatus", "Choice", "新規/対応中/保留/エスカレーション/解決済/クローズ（必須）"),
    ("優先度", "new_priority", "Choice", "高/中/低（必須）"),
    ("重要度", "new_severity", "Choice", "致命的/重大/軽微/影響なし"),
    ("カテゴリ", "new_categoryid", "Lookup", "→ カテゴリテーブル"),
    ("申請者", "new_requestorid", "Lookup", "→ SystemUser（必須）"),
    ("申請日時", "new_requestedon", "日時", "問い合わせ受付日時（必須）"),
    ("担当者", "crf98_personincharge", "Lookup", "→ SystemUser"),
    ("エスカレーション先", "new_escalatedtoid", "Lookup", "→ SystemUser"),
    ("対応期限", "new_duedate", "日時", "対応完了の期限"),
    ("質問内容", "new_questioncontent", "複数行テキスト", "問い合わせの詳細"),
    ("AI 回答案", "new_aisuggestion", "複数行テキスト", "AI が生成した回答案"),
    ("回答内容", "new_resolution", "複数行テキスト", "担当者の回答"),
    ("関連ナレッジ記事", "new_relatedknowledgearticleid", "Lookup", "→ ナレッジ記事"),
    ("解決日時", "new_resolvedon", "日時", "解決した日時"),
    ("クローズ日時", "new_closedon", "日時", "クローズした日時"),
    ("申請者名（外部）", "new_requestorname", "テキスト", "外部申請者の名前"),
    ("申請者メール（外部）", "new_requestoremail", "テキスト", "外部申請者のメール"),
]

# カテゴリテーブル列定義
CATEGORY_COLUMNS = [
    ("カテゴリ名", "new_name", "テキスト", "カテゴリの名前（必須）"),
    ("親カテゴリ", "new_parentcategoryid", "Lookup", "→ カテゴリ（自己参照・階層化用）"),
    ("表示順", "new_sortorder", "整数", "一覧での並び順"),
    ("説明", "new_description", "複数行テキスト", "カテゴリの説明"),
]

# ナレッジ記事テーブル列定義
KNOWLEDGE_COLUMNS = [
    ("タイトル", "new_title", "テキスト", "記事のタイトル（必須）"),
    ("記事番号", "new_articlenumber", "テキスト", "管理番号（例: KB-001）"),
    ("記事ステータス", "new_articlestatus", "Choice", "下書き/レビュー中/承認済み/公開/非公開/アーカイブ（必須）"),
    ("公開範囲", "new_visibility", "Choice", "全社公開/部門内/限定公開（必須）"),
    ("カテゴリ", "new_categoryid", "Lookup", "→ カテゴリテーブル"),
    ("質問・課題", "new_question", "複数行テキスト", "対象の問題・課題"),
    ("回答・解決策", "new_answer", "複数行テキスト", "解決方法"),
    ("備考", "new_remarks", "複数行テキスト", "補足情報"),
    ("レビュー担当", "new_reviewerid", "Lookup", "→ SystemUser"),
    ("レビュー担当割当日", "new_reviewassigneddate", "日時", "レビュー担当を割り当てた日"),
    ("レビュー完了期限", "new_reviewdeadline", "日時", "レビュー完了の期限"),
    ("レビュー実施日", "new_reviewdate", "日時", "実際にレビューした日"),
    ("移管先記事", "new_transfertoid", "Lookup", "→ ナレッジ記事（自己参照）"),
]

# フロー一覧（Appendix用）
FLOW_DETAILS = [
    ("チケットが起票された場合に呼び出しされる", "クラウドフロー", "チケット作成時", "担当者への通知メール送信"),
    ("チケットがクローズした場合にナレッジ化", "クラウドフロー", "チケットクローズ時", "ナレッジ記事の自動生成"),
    ("エージェントフロー_チケットを起票", "クラウドフロー", "エージェント呼出", "エージェント経由でチケット自動作成"),
    ("チケット管理プロセス", "BPF", "—", "チケットのライフサイクル管理"),
    ("ナレッジ記事管理プロセス", "BPF", "—", "記事の作成〜公開プロセス"),
]

# エージェント一覧
AGENTS = [
    ("Ask Copilot", "copilots_header_new_agent2", "担当者が回答案を作成するために利用"),
    ("Ask Specialist", "copilots_header_new_agent1", "エンドユーザーがチケットを起票"),
]

# チャート一覧
CHARTS = [
    ("ステータス別チケット件数", "横棒グラフ", "new_ticketstatus でグループ化"),
    ("優先度別チケット件数", "棒グラフ", "new_priority でグループ化"),
    ("カテゴリ別チケット件数", "横棒グラフ", "new_categoryid でグループ化"),
    ("チケット起票推移（月別）", "折れ線グラフ", "new_requestedon を月別集計"),
    ("チケット起票推移（日別）", "折れ線グラフ", "new_requestedon を日別集計"),
    ("担当者別チケット件数", "横棒グラフ", "crf98_personincharge でグループ化"),
    ("担当者×ステータス別", "積み上げ横棒", "担当者＋ステータスでクロス集計"),
    ("未対応チケット（担当者別）", "横棒グラフ", "ステータス1-4のみ、担当者でグループ化"),
    ("未対応チケット（ステータス×担当者）", "積み上げ横棒", "ステータス1-4のみ、担当者＋ステータスでクロス集計"),
]

# 免責事項
DISCLAIMER = """本アプリ集は日本マイクロソフトが提供する無償のサンプル群です。本アプリ集をダウンロードされた方は、以下の免責事項を承諾したものとみなされます。

1.\u3000本アプリ集（本アプリ集に付属するドキュメント及びReadmeに記載されている技術情報を含みます。以下、本「免責事項」において同じ。）は利用者に対して「現状のまま」提供されるものであり、日本マイクロソフトは、本アプリ集にプログラミング上の誤りその他の瑕疵のないこと、本アプリ集が利用者の目的に適合すること、並びに本アプリ集及びその使用が利用者または利用者以外の第三者の権利を侵害するものでないこと、その他のいかなる内容についての明示または黙示の保証を行うものではありません。

2.\u3000日本マイクロソフトは、本アプリ集の使用に起因して、利用者に生じた損害または第三者からの請求に基づく利用者の損害について、原因の如何を問わず、一切の責任を負いません。日本マイクロソフトは、本アプリ集に関連して利用者と第三者との間に発生するいかなる紛争について、一切責任を負わないものとします。本アプリ集の利用は、利用者の責任のもとで行ってください。

3.\u3000日本マイクロソフトは、本アプリ集の全部または一部の提供を廃止することがあります。提供の廃止によって利用者に発生した損害について、日本マイクロソフトは一切責任を負いません。

4.\u3000日本マイクロソフトは、本アプリ集のバグ修正、補修、保守、機能追加その他のいかなる義務も負いません。本アプリ集は、不定期に更新される可能性がありますが、バグ修正等が保証されているわけではありません。本アプリ集の安定した動作を確保するためには、利用者自身が適切なテストや検証を行ってください。

5.\u3000日本マイクロソフトは、本アプリ集に関するお問い合わせにはお答えできません。ご利用にあたっては、提供された手順書を参照し、ご自身でのインストールや利用を行ってください。"""


# ============================================================
# ヘルパー関数
# ============================================================
def add_table(doc, headers, rows):
    """ヘッダー行 + データ行のテーブルを追加"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            table.rows[ri + 1].cells[ci].text = str(val)
    return table


def add_screenshot_placeholder(doc, label):
    """スクリーンショットプレースホルダー"""
    p = doc.add_paragraph(f"【 {label} のスクリーンショットを挿入 】")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.italic = True


def add_tip(doc, text):
    """💡 補足テキスト"""
    doc.add_paragraph(f"💡 {text}")


# ============================================================
# ドキュメント生成
# ============================================================
doc = Document()

# ============================================================
# 表紙
# ============================================================
for _ in range(6):
    doc.add_paragraph("")
p = doc.add_paragraph("【サンプルアプリ】")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(APP_NAME)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(28)
    run.bold = True
p = doc.add_paragraph("導入手順書")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(18)
doc.add_paragraph("")
p = doc.add_paragraph(f"バージョン {SOLUTION_VERSION}\u3000｜\u3000{DOC_DATE}")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ============================================================
# 目次
# ============================================================
doc.add_heading("目次", level=1)
toc_items = [
    "1. アプリ概要",
    "2. インストール手順",
    "3. 追加設定",
    "4. 利用方法",
    "5. Appendix",
]
for item in toc_items:
    doc.add_paragraph(item)
doc.add_page_break()

# ============================================================
# 1. アプリ概要
# ============================================================
doc.add_heading("1. アプリ概要", level=1)
doc.add_paragraph(
    "汎用ヘルプデスクは、社内の問い合わせ（チケット）の受付・対応・ナレッジ管理を"
    "一元化する Copilot Studio のエージェントとモデル駆動型アプリです。"
)
doc.add_paragraph(
    "Copilot Studio エージェントによる AI 回答案の生成、チケットクローズ時の"
    "ナレッジ記事自動作成など、対応業務を効率化します。9 種類のチャートで"
    "チケットの状況をリアルタイムに可視化でき、ヘルプデスク業務の改善に活用できます。"
)

doc.add_heading("1.1 主な機能", level=2)
features = [
    "チケットの起票・ステータス管理・担当者割当",
    "対応履歴の記録（質問内容 → AI 回答案 → 回答内容）",
    "Copilot Studio エージェントによる AI 回答案生成",
    "ナレッジ記事の作成・レビュー・公開管理",
    "カテゴリによるチケット・記事の分類（階層対応）",
    "9 種類のチャートによる状況可視化",
    "ビジネスプロセスフロー（BPF）によるプロセス標準化",
    "チケットクローズ時のナレッジ記事自動生成",
    "チケット起票時の自動通知",
]
for f in features:
    doc.add_paragraph(f, style="List Bullet")

doc.add_heading("1.2 ソリューション構成", level=2)
add_table(doc, ["コンポーネント", "数量", "内容"], COMPONENTS)

doc.add_heading("1.3 必要な条件", level=2)
add_table(doc, ["要件", "詳細"], PREREQUISITES)

# ============================================================
# 2. インストール手順（★ 固定テンプレート）
# ============================================================
doc.add_heading("2. インストール手順", level=1)
doc.add_paragraph(
    "ソリューションファイル（ZIP）を Power Apps 環境にインポートします。"
    "以下の手順に従ってください。"
)

# 2.1
doc.add_heading("2.1 Power Apps ポータルにサインイン", level=2)
for step in [
    "ブラウザで https://make.powerapps.com/ にアクセスします",
    "管理者アカウントでサインインします",
]:
    doc.add_paragraph(step, style="List Number")

# 2.2
doc.add_heading("2.2 環境の選択", level=2)
for step in [
    "画面右上の環境セレクターをクリックします",
    "ソリューションをインポートする対象環境を選択します",
]:
    doc.add_paragraph(step, style="List Number")
add_screenshot_placeholder(doc, "環境セレクター")
add_tip(doc, "対象環境で Dataverse が有効になっていることを確認してください。")

# 2.3
doc.add_heading("2.3 ソリューションのインポート", level=2)
for step in [
    "左メニューの「ソリューション」をクリックします",
    "上部メニューの「インポート」をクリックします",
    "「ソリューションをインポート」ダイアログが表示されます",
    "「参照」ボタンをクリックし、ダウンロードしたソリューション ZIP ファイルを選択します",
    "「次へ」をクリックします",
]:
    doc.add_paragraph(step, style="List Number")
add_screenshot_placeholder(doc, "ソリューションインポート")

# 2.4
doc.add_heading("2.4 接続の設定", level=2)
doc.add_paragraph(
    "インポート中に接続の設定を求められる場合は、接続に対してアカウントを"
    "選択または新規作成してください。すべての接続を設定したら「次へ」を"
    "クリックします",
    style="List Number",
)
add_screenshot_placeholder(doc, "接続設定画面")
add_table(
    doc,
    ["接続", "用途", "設定方法"],
    [(c[0], c[1], "既存の接続を選択、またはサインインして新規作成") for c in CONNECTIONS],
)
doc.add_paragraph("")
doc.add_paragraph(
    "Dataverse URLの環境変数の設定を求められる場合は自環境の URL"
    "（例: https://xxxxx.crm.dynamics.com）を設定してください。"
    "アプリIDは固有のため、変更は不要です。",
    style="List Number",
)
add_tip(doc,
    "Dataverse URL は Power Platform 管理センター"
    "（https://admin.powerplatform.microsoft.com/）で確認できます。"
    "「環境」→ 対象環境を選択 →「環境 URL」に表示されています。"
)

# 2.5
doc.add_heading("2.5 インポートの実行", level=2)
for step in [
    "設定内容を確認し、「インポート」をクリックします",
    "インポートが開始されます。完了まで数分〜10 分程度かかります",
    f"ソリューション「{APP_NAME}」が正常にインポートされました、と表示されたら完了です",
]:
    doc.add_paragraph(step, style="List Number")
add_tip(doc, "インポートに失敗した場合は、エラーメッセージを確認してください。接続の認証エラーが最も多い原因です。")

# ============================================================
# 3. 追加設定
# ============================================================
doc.add_heading("3. 追加設定", level=1)
doc.add_paragraph("ソリューションインポート後、以下の設定を行ってください。")

# 3.1
doc.add_heading("3.1 Power Automate フローの有効化", level=2)
doc.add_paragraph("インポート後、フローが無効（Draft）状態になっている場合があります。")
for step in [
    "https://make.powerautomate.com/ にサインインします",
    f"左メニューの「ソリューション」→「{APP_NAME}」を開きます",
    "以下のフローを確認し、無効の場合は「有効にする」をクリックします",
]:
    doc.add_paragraph(step, style="List Number")
add_table(doc, ["フロー名", "動作", "確認事項"], FLOWS)
add_screenshot_placeholder(doc, "フローの有効化")

doc.add_heading("フローの接続修正", level=3)
doc.add_paragraph("フローが「接続エラー」で失敗する場合は、以下の手順で接続を修正します。")
for step in [
    "対象フローを開き「編集」をクリックします",
    "赤い警告が表示されている接続をクリックします",
    "「接続の修正」で正しいアカウントを選択します",
    "「保存」をクリックします",
]:
    doc.add_paragraph(step, style="List Number")
add_screenshot_placeholder(doc, "接続エラーの修正")

doc.add_heading("通知先メールアドレスの設定", level=3)
for step in [
    "「チケットが起票された場合に呼び出しされる」フローを編集します",
    "メール送信アクションの宛先を、組織のヘルプデスク担当者メールアドレスに変更します",
    "保存して有効化します",
]:
    doc.add_paragraph(step, style="List Number")
add_tip(doc, "メーリングリスト（例: helpdesk@yourcompany.com）を指定すると運用しやすいです。")

# 3.2
doc.add_heading("3.2 セキュリティロールの割当", level=2)
for step in [
    "Power Platform 管理センター > 環境 > 対象環境 > 設定 > ユーザー＋アクセス許可 > セキュリティロール を開きます",
    "「Basic User」ロールがアプリに関連付けられていることを確認します",
    "アプリ利用者にセキュリティロールを割り当てます",
]:
    doc.add_paragraph(step, style="List Number")
add_tip(doc, "セキュリティグループ単位で一括割当すると効率的です。")
add_screenshot_placeholder(doc, "セキュリティロールの割当")

# 3.3
doc.add_heading("3.3 カテゴリの初期設定", level=2)
doc.add_paragraph("カテゴリはソリューションに含まれないため、環境ごとに登録が必要です。")
for step in [
    "アプリを開き、左ナビの「カテゴリ」をクリックします",
    "「＋ 新規」をクリックして親カテゴリを作成します",
    "続けて子カテゴリを作成し、「親カテゴリ」フィールドで親を選択します",
]:
    doc.add_paragraph(step, style="List Number")
doc.add_paragraph("推奨カテゴリ構成:")
add_table(doc, ["親カテゴリ", "子カテゴリ"], CATEGORIES)

# 3.4
doc.add_heading("3.4 アプリの共有", level=2)
for step in [
    "make.powerapps.com > アプリ一覧を開きます",
    f"「{APP_NAME}」の「…」メニュー > 「共有」をクリックします",
    "利用者（ユーザーまたはセキュリティグループ）を追加します",
    "適切なセキュリティロールを選択します",
]:
    doc.add_paragraph(step, style="List Number")
add_screenshot_placeholder(doc, "アプリの共有画面")

# 3.5
doc.add_heading("3.5 Copilot Studio エージェントの設定（オプション）", level=2)
doc.add_paragraph(
    "AI 回答案生成・チケット自動起票機能を使用する場合は、"
    "Copilot Studio でエージェントの設定が必要です。"
)
for step in [
    "Copilot Studio（https://copilotstudio.microsoft.com/）にサインインします",
    "対象環境を選択し、「Ask Copilot」「Ask Specialist」エージェントを確認します",
    "必要に応じてナレッジソース・ツール・トリガーを設定します",
]:
    doc.add_paragraph(step, style="List Number")
add_tip(doc, "エージェントの詳細設定は Copilot Studio UI で行ってください。")

# ============================================================
# 4. 利用方法
# ============================================================
doc.add_heading("4. 利用方法", level=1)

# 4.1
doc.add_heading("4.1 アプリへのアクセス", level=2)
for step in [
    "ブラウザで Power Apps ポータル（https://make.powerapps.com/）を開きます",
    "左メニューの「アプリ」をクリックします",
    f"「{APP_NAME}」を選択して起動します",
]:
    doc.add_paragraph(step, style="List Number")

# 4.2
doc.add_heading("4.2 ナビゲーション", level=2)
doc.add_paragraph("アプリの左サイドバーには 3 つのグループがあります。")
add_table(doc, ["グループ", "メニュー", "説明"], NAVIGATION)

# 4.3
doc.add_heading("4.3 チケットの新規作成", level=2)
for step in [
    "左ナビの「チケット」をクリックします",
    "「＋ 新規」ボタンをクリックします",
    "必須項目を入力します: タイトル、優先度（高/中/低）、申請日時、申請者",
    "任意項目を入力します: カテゴリ、重要度、担当者、対応期限、質問内容",
    "「保存」をクリックします（チケット番号が自動採番されます）",
]:
    doc.add_paragraph(step, style="List Number")

# 4.4
doc.add_heading("4.4 チケットのステータス管理", level=2)
doc.add_paragraph("チケットには以下のステータスがあります。")
add_table(doc, ["ステータス", "説明", "次のステータス"], STATUSES)

# 4.5
doc.add_heading("4.5 チケットフォームの構成", level=2)
doc.add_paragraph("チケットの詳細画面には 3 つのタブがあります。")

doc.add_paragraph("「一般」タブ")
for item in [
    "基本情報: タイトル、チケット番号、ステータス、優先度、重要度、カテゴリ",
    "担当・期限: 申請者、申請日時、担当者、エスカレーション先、対応期限",
    "質問内容: 問い合わせの詳細テキスト",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph("「対応」タブ")
for item in [
    "AI 回答案: AI が生成した回答案（参考情報）",
    "回答内容: 担当者が記入する回答",
    "関連ナレッジ記事: 参考にしたナレッジ記事へのリンク",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph("「外部申請」タブ")
for item in [
    "申請者名（外部）、申請者メール（外部）: 社外からの問い合わせ用",
]:
    doc.add_paragraph(item, style="List Bullet")

add_screenshot_placeholder(doc, "チケットフォーム（対応タブ）")

# 4.6
doc.add_heading("4.6 チャートで状況を可視化", level=2)
doc.add_paragraph(
    "チケット一覧画面の右上のグラフアイコン（📊）をクリックすると、"
    "チャートパネルが表示されます。"
)
add_tip(doc, "チャートはビューのフィルタと連動します。ビューを切り替えるとチャートもそのデータで描画されます。")

# 4.7
doc.add_heading("4.7 ナレッジ記事の管理", level=2)
for step in [
    "左ナビの「ナレッジ記事」をクリックします",
    "「＋ 新規」で記事を作成します",
    "タイトル、記事番号、ステータス、公開範囲、カテゴリ、質問・課題、回答・解決策を入力します",
    "「保存」をクリックします",
]:
    doc.add_paragraph(step, style="List Number")
doc.add_paragraph("記事ステータスの流れ: 下書き → レビュー中 → 承認済み → 公開")

# 4.8
doc.add_heading("4.8 カテゴリの管理", level=2)
for step in [
    "左ナビの「カテゴリ」をクリックします",
    "「＋ 新規」でカテゴリを作成します",
    "親カテゴリを指定すれば階層化も可能です",
]:
    doc.add_paragraph(step, style="List Number")

# 4.9
doc.add_heading("4.9 Copilot Studio エージェントからの問い合わせ・チケット起票", level=2)
doc.add_paragraph(
    "本ソリューションには、Copilot Studio エージェントを通じて"
    "問い合わせやチケット起票を行う機能が含まれています。"
)

doc.add_heading("4.9.1 エージェント概要", level=3)
doc.add_paragraph("2 つのエージェントが用意されています。")
for item in [
    "Ask Specialist：エンドユーザーが問い合わせを行い、チケットを自動起票するエージェント",
    "Ask Copilot：ヘルプデスク担当者が AI の支援を受けて回答案を作成するエージェント",
]:
    doc.add_paragraph(item, style="List Bullet")
add_screenshot_placeholder(doc, "Copilot Studio エージェント一覧")

doc.add_heading("4.9.2 Ask Specialist（エンドユーザー向け）", level=3)
doc.add_paragraph(
    "エンドユーザーが問い合わせを入力すると、エージェントが内容を分析し、"
    "必要に応じてチケットを自動で起票します。"
)
for step in [
    "Copilot Studio またはTeams から Ask Specialist エージェントにアクセスします",
    "問い合わせ内容を自然言語で入力します（例:「VPN に接続できません」）",
    "エージェントがナレッジベースを検索し、回答を提示します",
    "回答で解決しない場合、エージェントがチケットの起票を提案します",
    "ユーザーが承認すると、チケットが自動で作成されます",
]:
    doc.add_paragraph(step, style="List Number")
add_screenshot_placeholder(doc, "Ask Specialist との会話画面")
add_screenshot_placeholder(doc, "Ask Specialist チケット起票の確認")

doc.add_heading("4.9.3 Ask Copilot（担当者向け）", level=3)
doc.add_paragraph(
    "ヘルプデスク担当者が、チケットの内容を基に AI の支援を受けて"
    "回答案を作成できます。"
)
for step in [
    "Copilot Studio または Teams から Ask Copilot エージェントにアクセスします",
    "チケットの質問内容をコピー＆ペースト、またはチケット番号を伝えます",
    "エージェントがナレッジベースと過去のチケット情報を基に回答案を生成します",
    "生成された回答案をチケットの「回答内容」フィールドにコピーして使用します",
]:
    doc.add_paragraph(step, style="List Number")
add_screenshot_placeholder(doc, "Ask Copilot との会話画面")
add_screenshot_placeholder(doc, "Ask Copilot 回答案の生成")

doc.add_heading("4.9.4 エージェントの利用チャネル", level=3)
doc.add_paragraph("エージェントは以下のチャネルから利用できます。")
for item in [
    "Copilot Studio テストチャット（開発・テスト用）",
    "Microsoft Teams（組織内展開に推奨）",
    "Web チャット（社内ポータルに埋め込み）",
]:
    doc.add_paragraph(item, style="List Bullet")
add_tip(doc, "Teams への公開は Copilot Studio の「チャネル」設定から行います。詳細は「3.5 Copilot Studio エージェントの設定」を参照してください。")
add_screenshot_placeholder(doc, "Teams でのエージェント利用画面")

# ============================================================
# 5. Appendix
# ============================================================
doc.add_heading("5. Appendix", level=1)

doc.add_heading("5.1 テーブル構成", level=2)

doc.add_heading("チケット（new_ticket）", level=3)
add_table(doc, ["列名", "論理名", "型", "説明"], TICKET_COLUMNS)

doc.add_heading("カテゴリ（new_category）", level=3)
add_table(doc, ["列名", "論理名", "型", "説明"], CATEGORY_COLUMNS)

doc.add_heading("ナレッジ記事（new_knowledgearticle）", level=3)
add_table(doc, ["列名", "論理名", "型", "説明"], KNOWLEDGE_COLUMNS)

doc.add_heading("5.2 Power Automate フロー一覧", level=2)
add_table(doc, ["フロー名", "種別", "トリガー", "動作"], FLOW_DETAILS)

doc.add_heading("5.3 Copilot Studio エージェント一覧", level=2)
add_table(doc, ["エージェント名", "スキーマ名", "用途"], AGENTS)

doc.add_heading("5.4 チャート一覧", level=2)
add_table(doc, ["チャート名", "種類", "データソース"], CHARTS)

# ============================================================
# 保存
# ============================================================
os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT_FILE)
print(f"✅ {OUTPUT_FILE} を生成しました")
