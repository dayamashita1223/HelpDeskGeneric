"""
手順書を再生成し、赤枠アノテーション付き画像を挿入する統合スクリプト
"""
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw
import os, tempfile, shutil

DOCS_DIR = r"C:\dev\CustomerZeroKnowledge\CodeAppsDevelopmentStandard\docs"
V2_PATH = os.path.join(DOCS_DIR, "汎用ヘルプデスク_ソリューション導入手順書_v2.docx")
NEW_PATH = os.path.join(DOCS_DIR, "汎用ヘルプデスク_ソリューション導入手順書_new.docx")
FINAL_PATH = os.path.join(DOCS_DIR, "汎用ヘルプデスク_ソリューション導入手順書.docx")

NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS_WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"

# ============================================================
# Step 1: Extract images from v2 and bake red annotations
# ============================================================
print("Step 1: Extracting and annotating images from v2...")
v2_doc = Document(V2_PATH)
tmp_dir = tempfile.mkdtemp(prefix="annotated_")

# Extract all images
image_blobs = {}
for rel_id, rel in v2_doc.part.rels.items():
    if "image" in rel.reltype:
        img_name = os.path.basename(rel.target_ref)
        img_path = os.path.join(tmp_dir, img_name)
        with open(img_path, "wb") as f:
            f.write(rel.target_part.blob)
        image_blobs[rel_id] = img_path

# Find anchor shapes (red rects) and their associated inline images
# in paragraphs 45, 51, 54 of v2
annotated = {}
for pi in [45, 51, 54]:
    para = v2_doc.paragraphs[pi]

    # Get inline image info
    inline_imgs = para._element.findall(".//{%s}inline" % NS_WP)
    img_rel_id = None
    img_cx = img_cy = 0
    for inl in inline_imgs:
        ext = inl.find("{%s}extent" % NS_WP)
        if ext is not None:
            img_cx = int(ext.get("cx", 0))
            img_cy = int(ext.get("cy", 0))
        for blip in inl.findall(".//{%s}blip" % NS_A):
            img_rel_id = blip.get("{%s}embed" % NS_R)

    if not img_rel_id:
        continue

    # Get red rectangle positions
    rects = []
    for anchor in para._element.findall(".//{%s}anchor" % NS_WP):
        wsp = anchor.find(".//{%s}wsp" % NS_WPS)
        if wsp is None:
            continue
        ln = wsp.find(".//{%s}ln" % NS_A)
        if ln is None:
            continue
        clr = ln.find(".//{%s}srgbClr" % NS_A)
        if clr is None or clr.get("val") != "C00000":
            continue

        posH = anchor.find("{%s}positionH" % NS_WP)
        posV = anchor.find("{%s}positionV" % NS_WP)
        off_x = off_y = 0
        if posH is not None:
            po = posH.find("{%s}posOffset" % NS_WP)
            if po is not None and po.text:
                off_x = int(po.text)
        if posV is not None:
            po = posV.find("{%s}posOffset" % NS_WP)
            if po is not None and po.text:
                off_y = int(po.text)
        ext = anchor.find("{%s}extent" % NS_WP)
        cx = int(ext.get("cx", 0)) if ext is not None else 0
        cy = int(ext.get("cy", 0)) if ext is not None else 0
        rects.append((off_x, off_y, cx, cy))

    annotated[img_rel_id] = {"img_cx": img_cx, "img_cy": img_cy, "rects": rects}

# Bake red rectangles into images
for rel_id, info in annotated.items():
    img_path = image_blobs[rel_id]
    img = Image.open(img_path)
    img_w, img_h = img.size
    scale_x = img_w / info["img_cx"]
    scale_y = img_h / info["img_cy"]
    draw = ImageDraw.Draw(img)
    for ox, oy, cx, cy in info["rects"]:
        px_x = int(ox * scale_x)
        px_y = int(oy * scale_y)
        px_w = int(cx * scale_x)
        px_h = int(cy * scale_y)
        line_w = max(3, int(3 * scale_x))
        draw.rectangle([px_x, px_y, px_x + px_w, px_y + px_h], outline="#C00000", width=line_w)
    img.save(img_path, "PNG")
    print(f"  Annotated {rel_id} ({os.path.basename(img_path)}): {len(info['rects'])} rects")

print(f"  {len(image_blobs)} images ready")

# ============================================================
# Step 2: Regenerate fresh docx
# ============================================================
print("\nStep 2: Regenerating deployment guide...")
gen_script = os.path.join(os.path.dirname(__file__), "generate_deployment_guide.py")
# Override OUTPUT_FILE before exec
import types
gen_code = open(gen_script, encoding="utf-8").read()
gen_code = gen_code.replace(
    'OUTPUT_FILE = os.path.join(OUTPUT_DIR, f"{APP_NAME}_ソリューション導入手順書.docx")',
    f'OUTPUT_FILE = r"{NEW_PATH}"',
)
exec(gen_code)

# ============================================================
# Step 3: Insert annotated images into new docx
# ============================================================
print("\nStep 3: Inserting images into guide...")
doc = Document(NEW_PATH)
count = 0


def rep_placeholder(doc, txt, img_path):
    """Replace placeholder text with image."""
    for p in doc.paragraphs:
        if txt in p.text:
            p.clear()
            r = p.add_run()
            r.add_picture(img_path, width=Inches(5.5))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
    return False


def add_img_after(doc, search_text, img_path):
    """Add image paragraph after paragraph containing search_text."""
    for p in doc.paragraphs:
        if search_text in p.text:
            np = doc.add_paragraph()
            np.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = np.add_run()
            r.add_picture(img_path, width=Inches(5.5))
            p._element.addnext(np._element)
            return True
    return False


# Section 2 images (with red annotations)
if rep_placeholder(doc, "環境セレクター のスクリーンショットを挿入", image_blobs["rId8"]):
    count += 1; print("  ✓ 2.2 環境セレクター (annotated)")
if rep_placeholder(doc, "ソリューションインポート のスクリーンショットを挿入", image_blobs["rId9"]):
    count += 1; print("  ✓ 2.3 ソリューションインポート (annotated)")
if add_img_after(doc, "「次へ」をクリックします", image_blobs["rId10"]):
    count += 1; print("  ✓ 2.3 ZIP選択画面 (annotated)")
if rep_placeholder(doc, "接続設定画面 のスクリーンショットを挿入", image_blobs["rId11"]):
    count += 1; print("  ✓ 2.4 接続設定画面")
if add_img_after(doc, "アプリIDは固有のため、変更は不要です", image_blobs["rId12"]):
    count += 1; print("  ✓ 2.4 環境変数設定画面")

# Section 4 images (no annotations)
if add_img_after(doc, "を選択して起動します", image_blobs["rId13"]):
    count += 1; print("  ✓ 4.1 アプリ起動画面")
if add_img_after(doc, "3 つのグループがあります", image_blobs["rId14"]):
    count += 1; print("  ✓ 4.2 ナビゲーション")
if add_img_after(doc, "チケット番号が自動採番されます", image_blobs["rId15"]):
    count += 1; print("  ✓ 4.3 チケット新規作成")
if add_img_after(doc, "以下のステータスがあります", image_blobs["rId16"]):
    count += 1; print("  ✓ 4.4 ステータス管理")
if add_img_after(doc, "チャートパネルが表示されます", image_blobs["rId17"]):
    count += 1; print("  ✓ 4.6 チャート")
if add_img_after(doc, "記事ステータスの流れ", image_blobs["rId18"]):
    count += 1; print("  ✓ 4.7 ナレッジ記事")
if add_img_after(doc, "階層化も可能です", image_blobs["rId19"]):
    count += 1; print("  ✓ 4.8 カテゴリ")

doc.save(NEW_PATH)
shutil.rmtree(tmp_dir)
print(f"\n✅ {count}/12 images inserted. Saved to {NEW_PATH}")
